import docx
from docx.shared import Inches
import os

def add_workflow_to_word():
    base_path = r'c:\Users\jeffj\OneDrive\Documents\laravel\book_vault'
    doc_path = os.path.join(base_path, 'BookVault_Project_Report.docx')
    image_path = os.path.join(base_path, 'diagrams', 'workflow_premium.png')
    output_path = os.path.join(base_path, 'BookVault_Project_Report_Final.docx')

    if not os.path.exists(doc_path):
        print(f"Error: {doc_path} not found.")
        return
    if not os.path.exists(image_path):
        print(f"Error: {image_path} not found.")
        return

    doc = docx.Document(doc_path)
    
    # Find where to insert (under "Workflow Diagram" or near section 4.3)
    # The Markdown says 63. - Workflow Diagram
    # I'll look for "Workflow Diagram"
    
    insert_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "Workflow Diagram" in p.text:
            insert_idx = i
            break
            
    if insert_idx == -1:
        # Fallback to Architecture Overview
        for i, p in enumerate(doc.paragraphs):
            if "Architecture Overview" in p.text:
                insert_idx = i
                break

    if insert_idx != -1:
        target = doc.paragraphs[insert_idx]
        
        # Add the image after the paragraph
        # doc.add_picture appends to the end. We need to insert.
        run = target.add_run()
        run.add_break()
        run.add_picture(image_path, width=Inches(6))
        
        print(f"Successfully added workflow diagram to {output_path}")
    else:
        print("Could not find insertion point.")

    doc.save(output_path)

if __name__ == '__main__':
    add_workflow_to_word()
