import os
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_mermaid_image(mermaid_code, filename):
    url = 'https://kroki.io/mermaid/png'
    try:
        response = requests.post(url, data=mermaid_code.encode('utf-8'), headers={'Content-Type': 'text/plain'}, timeout=15)
        if response.status_code == 200:
            with open(filename, 'wb') as f:
                f.write(response.content)
            print(f"Downloaded {filename}")
        else:
            print(f"Failed to download {filename}: {response.status_code} - {response.text[:100]}")
    except Exception as e:
        print(f"Error downloading {filename}: {e}")

diagrams = {
    'uc': '''
flowchart LR
    User -->|Views Collection| Book[Book System]
    User -->|Borrows Book| Lending[Lending System]
    User -->|Buys/Sells Book| Marketplace[Marketplace System]
    Librarian -->|Moderates Books| Book
    Librarian -->|Moderates Marketplace| Marketplace
    Admin -->|Manages Users & Roles| UserSystem[User System]
    ''',
    
    'seq_login': '''
sequenceDiagram
    participant User
    participant System
    participant Database
    User->>System: Enter Email & Password
    System->>Database: Validate Credentials
    Database-->>System: Validation Result
    alt Valid
        System-->>User: Redirect to Dashboard
    else Invalid
        System-->>User: Show Error Message
    end
    ''',
    
    'seq_lend': '''
sequenceDiagram
    participant User
    participant System
    participant Book
    participant Lending
    User->>System: Request Borrow
    System->>Book: Check Availability
    alt Available
        System->>Lending: Create Record
        System->>Book: Update Stock
        System-->>User: Request Success
    else Unavailable
        System-->>User: Show Error
    end
    ''',
    
    'er': '''
erDiagram
    USER ||--o{ LENDING : borrows
    USER ||--o{ BOOKFORSALE : lists
    USER ||--o{ PURCHASE : makes
    BOOK ||--o{ LENDING : involves
    BOOKFORSALE ||--o{ PURCHASE : involves
    LENDING ||--o| FINE : incurs
    ''',

    'class': '''
classDiagram
    class User {
        +id: Integer
        +name: String
        +email: String
    }
    class Role {
        +id: Integer
        +name: String
    }
    class Book {
        +id: Integer
        +title: String
        +stock: Integer
    }
    class Lending {
        +id: Integer
        +user_id: Integer
        +book_id: Integer
        +start_date: Date
        +return_date: Date
    }
    class Fine {
        +id: Integer
        +amount: Double
        +paid: Boolean
    }
    class BookForSale {
        +id: Integer
        +title: String
        +price: Double
    }
    class Purchase {
        +id: Integer
        +user_id: Integer
        +book_for_sale_id: Integer
    }
    User "1" -- "many" Role : has
    User "1" -- "many" Lending : creates
    Book "1" -- "many" Lending : associated
    Lending "1" -- "0..1" Fine : triggers
    User "1" -- "many" BookForSale : lists 
    User "1" -- "many" Purchase : makes
    BookForSale "1" -- "many" Purchase : associated
    '''
}

for name, code in diagrams.items():
    get_mermaid_image(code, f"{name}.png")

doc = Document()

# Add college details to match the format somewhat
doc.add_heading('WPL Project Report', 0)
p = doc.add_paragraph('BookVault')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Done by Jeff Jacob Jaison').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Guided by Mr. Ajay Antony Joseph').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

doc.add_heading('ABSTRACT', level=1)
doc.add_paragraph("BookVault is a comprehensive web-based platform designed to simplify library management, integrating lending services, marketplace features, user interactions, and fines calculation. Built on Laravel 12, BookVault provides a dynamic environment for members to borrow books, submit e-books, and buy or sell personal books. Librarians and Administrators ensure a seamless ecosystem with role-based moderation. This modern project enhances operational efficiency in book distribution and digital cataloging.")
doc.add_page_break()

doc.add_heading('1. INTRODUCTION', level=1)
doc.add_paragraph("BookVault is a robust application developed using the latest Laravel framework. It provides users with an integrated collection of both library lending services and a peer-to-peer marketplace. Members can view available books, track their borrow history, resolve their fines, and participate in community sales. With careful validation and security layers, BookVault ensures high-level data integrity and scalable architecture in managing digital and physical books.")

doc.add_heading('2. SYSTEM ANALYSIS', level=1)
doc.add_heading('2.1 Module Description', level=2)
doc.add_paragraph("1. Book Management: Enables Librarians to moderate books, modify stocks, and categorize digital submissions.\n2. Lending Management: Users can request books and track return dates. The system deducts stock on lending and restores it upon return.\n3. Fine Management: Calculates penalties dynamically if a user exceeds the lending deadline.\n4. Marketplace: Allows members to list personal books for sale and others to purchase them.\n5. Role Management: Secures access through specific roles (Member, Librarian, Admin).")

doc.add_heading('2.2 Actors and Roles', level=2)
doc.add_paragraph("Admin: Has ultimate control over the platform's settings, users, and overall operations.\nLibrarian: Focuses on moderating book collections, approving e-book submissions, and managing the market listings.\nMember (User): The central consumer who borrows books, pays fines, lists items for sale, and makes purchases.")
doc.add_page_break()

doc.add_heading('3. SYSTEM DESIGN', level=1)
doc.add_heading('3.1 Use Case Model', level=2)
if os.path.exists('uc.png'):
    doc.add_picture('uc.png', width=Inches(5))
    
doc.add_heading('3.2 Sequence Diagrams', level=2)
doc.add_paragraph("User Login:")
if os.path.exists('seq_login.png'):
    doc.add_picture('seq_login.png', width=Inches(4.5))

doc.add_paragraph("Book Lending:")
if os.path.exists('seq_lend.png'):
    doc.add_picture('seq_lend.png', width=Inches(4.5))

doc.add_page_break()
doc.add_heading('3.3 Entity Relationship Diagram', level=2)
if os.path.exists('er.png'):
    doc.add_picture('er.png', width=Inches(5))
    
doc.add_heading('3.4 Class Diagram', level=2)
if os.path.exists('class.png'):
    doc.add_picture('class.png', width=Inches(5))

doc.add_page_break()
doc.add_heading('4. SYSTEM ARCHITECTURE & DATABASE DESIGN', level=1)
doc.add_paragraph("The application relies on a modern MVC architecture provided by the Laravel framework. The backend logic securely connects to the PostgreSQL/MySQL database through Laravel's highly optimized Eloquent ORM. The frontend leverages Tailwind CSS / Bootstrap along with Vue/Inertia.js for seamless view rendering without continuous page reloads.")
doc.add_heading('4.1 Table Design', level=2)
doc.add_paragraph("1. Users Table: Central identity tracker, housing fields like: ID, Name, Email, Password, Role_ID.\n2. Books Table: Core entity for library stock, containing: Title, ISBN, Author, Base Stock, Approved status.\n3. Lendings Table: Logs active or resolved borrows, featuring tracking tokens: User_ID, Book_ID, Borrow_Date, Return_Deadline.\n4. Fines Table: Calculated values linking specific Lendings to overdue amounts, storing boolean resolution states.\n5. BookForSale and Purchases: Facilitator of the integrated P2P sales mechanics, linking multiple Users to specific items and price tags.")

doc.add_page_break()
doc.add_heading('5. CONCLUSION', level=1)
doc.add_paragraph("BookVault streamlines advanced library and reading processes, making book lending, selling, and tracking much easier and far more efficient. Through the use of modern development paradigms built on the Laravel framework, the platform remains highly reliable, expandable, and accessible.")

doc.save('BookVault_Project_Report.docx')
print("Report successfully generated: BookVault_Project_Report.docx")
