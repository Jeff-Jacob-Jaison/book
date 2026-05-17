import sys
import os
import re
try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("python-docx not installed. Skipping docx formatting.")
    sys.exit(0)

md_file = r'c:\Users\jeffj\OneDrive\Documents\laravel\book_vault\BookVault_New_Project_Report.md'
docx_file = r'c:\Users\jeffj\OneDrive\Documents\laravel\book_vault\BookVault_New_Project_Report.docx'

with open(md_file, 'r', encoding='utf-8') as f:
    lines = f.readlines()

doc = Document()
for line in lines:
    line = line.strip()
    if not line:
        continue
    
    # Handle Headers
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    
    # Handle Images: ![Alt Text](relative/path/to/image.png)
    elif line.startswith('!['):
        match = re.search(r'!\[.*?\]\((.*?)\)', line)
        if match:
            image_path = match.group(1)
            # Handle absolute/relative paths
            if not os.path.isabs(image_path):
                base_dir = os.path.dirname(md_file)
                image_path = os.path.join(base_dir, image_path)
            
            if os.path.exists(image_path):
                doc.add_picture(image_path, width=Inches(6))
            else:
                doc.add_paragraph(f"[Image Missing: {image_path}]")
    
    # Handle Lists
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif re.match(r'^\d+\. ', line):
        doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
    
    # Handle Horizontal Rules
    elif line == '---' or line == '***':
        doc.add_page_break()
    
    # Normal Paragraphs
    else:
        # Simple cleanup for bold/italic markdown
        clean_text = line.replace('**', '').replace('*', '')
        doc.add_paragraph(clean_text)

doc.save(docx_file)
print(f"Successfully updated {docx_file}")
