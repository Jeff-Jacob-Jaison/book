import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import os

def replace_workflow_with_basic():
    base_path = r'c:\Users\jeffj\OneDrive\Documents\laravel\book_vault'
    doc_path = os.path.join(base_path, 'BookVault_Project_Report.docx')
    src_image = os.path.join(base_path, 'diagrams', 'workflow_basic.png')
    target_image = os.path.join(base_path, 'diagrams', 'workflow_basic_processed.png')
    output_path = os.path.join(base_path, 'BookVault_Project_Report_Final.docx')

    if not os.path.exists(doc_path):
        print(f"Error: {doc_path} not found.")
        return
    if not os.path.exists(src_image):
        print(f"Error: {src_image} not found.")
        return

    # Process image with Pillow
    try:
        with Image.open(src_image) as img:
            img.save(target_image, "PNG")
    except Exception as e:
        print(f"Image processing failed: {e}")
        return

    doc = docx.Document(doc_path)
    
    # We'll re-run the insertion logic but first remove the old "3.6 Workflow Diagram" section
    # Actually, it's easier to use the backup (which doesn't have 3.6) and re-add 3.6 correctly.
    # Backup6 was before I added 3.6.
    
    doc = docx.Document(os.path.join(base_path, 'BookVault_Project_Report_Backup6.docx'))
    
    insert_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "4. SYSTEM ARCHITECTURE" in p.text.upper():
            insert_idx = i
            break
            
    if insert_idx == -1: insert_idx = len(doc.paragraphs)
    target = doc.paragraphs[insert_idx]
    
    header = target.insert_paragraph_before("3.6 Workflow Diagram")
    header.style = 'Heading 2' if 'Heading 2' in [s.name for s in doc.styles] else None
    
    p_desc = target.insert_paragraph_before("The following diagram illustrates the integrated workflow across different user roles within the BookVault system.")
    
    # Add Basic Image
    run = target.insert_paragraph_before("").add_run()
    run.add_picture(target_image, width=Inches(6))
    
    caption = target.insert_paragraph_before("Figure 3.6: BookVault Integrated Workflow Diagram (Basic)")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs: run.italic = True
    
    doc.save(output_path)
    print(f"Successfully replaced workflow section with basic version in {output_path}")

if __name__ == '__main__':
    replace_workflow_with_basic()
