import docx
from docx.shared import Inches

def update_report():
    doc = docx.Document('BookVault_Project_Report.docx')
    
    # We will iterate through paragraphs and modify as needed
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        
        # Abstract update
        if "BookVault is a comprehensive web-based platform" in text and "fine calculation" in text:
            p.text = p.text.replace("fines calculation.", "fines calculation, and a new personalized wishlist feature.")
        elif "Librarians and Administrators ensure a seamless ecosystem" in text:
            p.text = p.text.replace("Librarians and Administrators ensure", "Librarians ensure")

        # 2.1 Module Description update
        if "5. Role Management:" in text:
            p.text = "5. Role Management: Secures access through specific roles (Member, Librarian)."
            # Insert wishlist module description right after
            new_p = doc.paragraphs[i].insert_paragraph_before("6. Wishlist: Allows members to save books for later preview and access.")
            doc.paragraphs[i].text, new_p.text = new_p.text, doc.paragraphs[i].text # Swap to place it after
            
        # 2.2 Actors and Roles update
        if "Admin: Has ultimate control over the platform's settings" in text:
            p.clear() # remove admin line
            
        if "Librarian: Focuses on moderating book collections" in text:
            p.text = "Librarian: Has ultimate control over the platform's settings, moderating book collections, approving e-book submissions, and managing market listings."
            
        # Diagrams
        if "3.1 Use Case Model" in text:
            # We'll just add the image to a new run
            run = p.add_run()
            run.add_break()
            run.add_picture('diagrams/use_case.png', width=Inches(6.0))
            
        if "Book Lending:" in text:
            # Add buying sequence after book lending
            new_p = doc.paragraphs[i].insert_paragraph_before("Buying Sequence:")
            doc.paragraphs[i].text, new_p.text = new_p.text, doc.paragraphs[i].text # Swap to place it after
            
            # The next paragraph will be the buying sequence header, so we add a run to it with the image
            run = doc.paragraphs[i].add_run()
            run.add_break()
            try:
                run.add_picture('diagrams/buying_sequence.png', width=Inches(6.0))
            except Exception as e:
                print(f"Skipping buying sequence image: {e}")

        # 4.1 Table Design update
        if "5. BookForSale and Purchases" in text:
            new_p = doc.paragraphs[i].insert_paragraph_before("6. Wishlists Table: Stores user bookmarked items.")
            doc.paragraphs[i].text, new_p.text = new_p.text, doc.paragraphs[i].text # Swap to place it after

        # Insert UI Screenshots right before "4. SYSTEM ARCHITECTURE & DATABASE DESIGN"
        if "4. SYSTEM ARCHITECTURE & DATABASE DESIGN" in text:
            ui_p = p.insert_paragraph_before("3.5 UI Previews")
            ui_p.style = 'Heading 2' if 'Heading' in [s.name for s in doc.styles] else p.style
            
            p.insert_paragraph_before("Collection Page:")
            cp = p.insert_paragraph_before("")
            try:
                cp.add_run().add_picture('diagrams/collection_page.png', width=Inches(6.0))
            except: pass
            
            p.insert_paragraph_before("Book Preview Modal:")
            mp = p.insert_paragraph_before("")
            try:
                mp.add_run().add_picture('diagrams/preview_modal.png', width=Inches(6.0))
            except: pass
            
            p.insert_paragraph_before("Wishlist Page:")
            wp = p.insert_paragraph_before("")
            try:
                wp.add_run().add_picture('diagrams/wishlist_page.png', width=Inches(6.0))
            except: pass

    doc.save('BookVault_Project_Report_Updated.docx')
    print("Successfully updated BookVault_Project_Report_Updated.docx")

if __name__ == '__main__':
    update_report()
